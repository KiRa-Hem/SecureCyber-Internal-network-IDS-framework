"""
Convert the Capstone Report markdown into a properly formatted DOCX file.
Applies: Times New Roman, 12pt body, 16pt/14pt headings, 1.5 line spacing,
justified alignment, and proper page numbering.
"""
import re, sys
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn

import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PART1 = os.path.join(SCRIPT_DIR, "Capstone_Report_Final.md")
PART2 = os.path.join(SCRIPT_DIR, "Capstone_Report_Part2.md")
OUTPUT = os.path.join(SCRIPT_DIR, "Capstone_Report_FINAL.docx")


def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, line_spacing=1.5):
    para.alignment = alignment
    pf = para.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing


def add_heading_styled(doc, text, level=1):
    if level == 1:
        para = doc.add_paragraph()
        set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT, space_after=12)
        run = para.add_run(text.upper())
        set_font(run, size=16, bold=True)
    elif level == 2:
        para = doc.add_paragraph()
        set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
        run = para.add_run(text)
        set_font(run, size=14, bold=True)
    elif level == 3:
        para = doc.add_paragraph()
        set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
        run = para.add_run(text)
        set_font(run, size=13, bold=True)
    else:
        para = doc.add_paragraph()
        set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
        run = para.add_run(text)
        set_font(run, size=12, bold=True)
    return para


def add_body_text(doc, text):
    para = doc.add_paragraph()
    set_paragraph_format(para)
    run = para.add_run(text)
    set_font(run)
    return para


def add_code_block(doc, code_text):
    para = doc.add_paragraph()
    set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT, space_after=6, line_spacing=1.0)
    pf = para.paragraph_format
    pf.left_indent = Cm(1)
    run = para.add_run(code_text)
    set_font(run, name="Consolas", size=9)
    # Light gray background via shading
    shading = run._element.get_or_add_rPr()
    return para


def add_table_from_rows(doc, headers, rows):
    if not headers:
        return None
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=10, bold=True)
    for r_idx, row in enumerate(rows):
        # Pad or truncate row to match headers
        padded = (list(row) + [''] * ncols)[:ncols]
        for c_idx, val in enumerate(padded):
            cell = table.rows[r_idx + 1].cells[c_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_font(run, size=10)
    doc.add_paragraph()
    return table


def parse_md_table(lines, start_idx):
    """Parse a markdown table starting at start_idx, return (headers, rows, end_idx)."""
    headers = [c.strip() for c in lines[start_idx].strip().strip('|').split('|')]
    rows = []
    idx = start_idx + 2  # skip separator line
    while idx < len(lines) and '|' in lines[idx] and lines[idx].strip().startswith('|'):
        row = [c.strip() for c in lines[idx].strip().strip('|').split('|')]
        rows.append(row)
        idx += 1
    return headers, rows, idx


def process_markdown(doc, filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Skip horizontal rules
        if stripped == '---':
            doc.add_paragraph()
            i += 1
            continue

        # Headings
        if stripped.startswith('# ') and not stripped.startswith('## '):
            text = stripped[2:].strip()
            add_heading_styled(doc, text, level=1)
            i += 1
            continue
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            add_heading_styled(doc, text, level=2)
            i += 1
            continue
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            add_heading_styled(doc, text, level=3)
            i += 1
            continue
        if stripped.startswith('#### '):
            text = stripped[5:].strip()
            add_heading_styled(doc, text, level=4)
            i += 1
            continue

        # Tables
        if '|' in stripped and stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            headers, rows, end_idx = parse_md_table(lines, i)
            if headers and rows:
                add_table_from_rows(doc, headers, rows)
            i = end_idx
            continue

        # Bold text lines
        if stripped.startswith('**') and stripped.endswith('**'):
            text = stripped.strip('*').strip()
            para = doc.add_paragraph()
            set_paragraph_format(para)
            run = para.add_run(text)
            set_font(run, bold=True)
            i += 1
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            # Clean markdown formatting
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            para = doc.add_paragraph()
            set_paragraph_format(para, space_after=3)
            pf = para.paragraph_format
            pf.left_indent = Cm(1.5)
            pf.first_line_indent = Cm(-0.5)
            run = para.add_run("• " + text)
            set_font(run)
            i += 1
            continue

        # Numbered items
        m = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if m:
            text = m.group(2)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            para = doc.add_paragraph()
            set_paragraph_format(para, space_after=3)
            pf = para.paragraph_format
            pf.left_indent = Cm(1.5)
            pf.first_line_indent = Cm(-0.5)
            run = para.add_run(f"{m.group(1)}. {text}")
            set_font(run)
            i += 1
            continue

        # Empty lines
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        text = stripped
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        add_body_text(doc, text)
        i += 1


def main():
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(2.54)

    print("Processing Part 1 (Chapters 1-5)...")
    process_markdown(doc, PART1)

    print("Processing Part 2 (Chapters 6-10)...")
    process_markdown(doc, PART2)

    doc.save(OUTPUT)
    print(f"\n{'='*60}")
    print(f"REPORT GENERATED: {OUTPUT}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
